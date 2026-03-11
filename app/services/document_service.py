# app/services/document_service.py
"""
Serviço de geração automática de documentos.
Usa python-docx para preencher diretamente as tabelas do template do condomínio.
Suporta análise inteligente de PDF para mapeamento automático de campos.
"""
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil
import io
import re
import json
import base64

from app.config import settings
from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentService:
    """
    Serviço para geração de documentos a partir do template do condomínio.

    O template possui 3 tabelas:
    - Table 0: Dados do proprietário (FIXOS, já preenchidos no template)
    - Table 1: Dados do hóspede (5 rows x 6 cols)
    - Table 2: Acompanhantes (7 rows x 3 cols)

    E um parágrafo com Veículo/Modelo e Placa.
    """

    def __init__(self):
        """Inicializa o serviço de documentos"""
        self.template_dir = Path(settings.TEMPLATE_DIR)
        self.output_dir = Path(settings.OUTPUT_DIR)

        # Garantir que diretórios existem
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _find_template(self) -> Optional[Path]:
        """
        Encontra o template do condomínio no diretório de templates.
        Procura pelo DEFAULT_TEMPLATE ou qualquer .docx no diretório.
        """
        # Tentar o template padrão configurado
        default_path = self.template_dir / settings.DEFAULT_TEMPLATE
        if default_path.exists():
            return default_path

        # Procurar qualquer .docx no diretório de templates
        docx_files = list(self.template_dir.glob("*.docx"))
        if docx_files:
            logger.info(f"Using template: {docx_files[0].name}")
            return docx_files[0]

        return None

    def _set_cell_text(self, table, row_idx: int, col_idx: int, text: str):
        """Define o texto de uma célula mantendo a formatação existente."""
        cell = table.rows[row_idx].cells[col_idx]
        # Limpar o conteúdo existente
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
        # Definir o novo texto no primeiro run ou criar um
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = text
        elif cell.paragraphs:
            cell.paragraphs[0].text = text
        else:
            cell.text = text

    def _insert_logo_header(self, doc, logo_url: str) -> None:
        """
        Insere a logo do condomínio no início do documento.
        Suporta base64 data URI e URL pública.
        Falha graciosamente — não interrompe a geração do documento se a logo não puder ser inserida.
        """
        try:
            if logo_url.startswith('data:'):
                # Base64 data URI: data:image/png;base64,<data>
                _, encoded = logo_url.split(',', 1)
                # Corrigir padding base64 se necessário
                padding = 4 - len(encoded) % 4
                if padding != 4:
                    encoded += '=' * padding
                img_io = io.BytesIO(base64.b64decode(encoded))
            else:
                # URL pública — download com timeout curto
                import requests as http_requests
                resp = http_requests.get(logo_url, timeout=5)
                resp.raise_for_status()
                img_io = io.BytesIO(resp.content)

            # Adicionar parágrafo com imagem (alinhado à direita)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run()
            run.add_picture(img_io, height=Cm(1.8))

            # Mover elemento XML para o início do body (antes de todas as tabelas)
            doc.element.body.insert(0, para._element)

        except Exception as e:
            logger.warning(f"[DocumentService] Logo insert failed (non-fatal): {e}")

    def generate_condo_authorization(
        self,
        booking_data: Dict[str, Any],
        property_data: Dict[str, Any],
        guest_data: Dict[str, Any],
        save_to_file: bool = True,
        logo_url: str = "",
    ) -> Dict[str, Any]:
        """
        Gera autorização de condomínio preenchendo o template real.

        Preenche:
        - Table 1 (hóspede): nome, CPF, endereço, telefone, bairro, celular,
          cidade, estado, entrada, saída, CEP
        - Table 2 (acompanhantes): nome e documento de cada acompanhante
        - Parágrafo de veículo/placa
        """
        try:
            template_path = self._find_template()

            if not template_path:
                logger.warning("No template found, creating default")
                template_path = self.template_dir / settings.DEFAULT_TEMPLATE
                self._create_default_template(template_path)

            # Copiar template para não modificar o original
            doc = Document(str(template_path))

            # Inserir logo do condomínio no início do documento (se configurada)
            if logo_url:
                self._insert_logo_header(doc, logo_url)

            # Formatar datas
            check_in = booking_data.get('check_in')
            check_out = booking_data.get('check_out')

            if isinstance(check_in, str):
                check_in_formatted = datetime.fromisoformat(check_in).strftime('%d/%m/%Y')
            else:
                check_in_formatted = check_in.strftime('%d/%m/%Y') if check_in else ''

            if isinstance(check_out, str):
                check_out_formatted = datetime.fromisoformat(check_out).strftime('%d/%m/%Y')
            else:
                check_out_formatted = check_out.strftime('%d/%m/%Y') if check_out else ''

            # Verificar se o documento tem tabelas suficientes
            if len(doc.tables) >= 2:
                # ========== TABLE 1: Dados do Hóspede ==========
                guest_table = doc.tables[1]

                # Row 0: Hóspede / CPF
                self._set_cell_text(guest_table, 0, 1, guest_data.get('name', ''))
                self._set_cell_text(guest_table, 0, 5, guest_data.get('cpf', guest_data.get('document_number', '')))

                # Row 1: Endereço / Telefone
                self._set_cell_text(guest_table, 1, 1, guest_data.get('address', ''))
                self._set_cell_text(guest_table, 1, 5, guest_data.get('phone', ''))

                # Row 2: Bairro / Celular
                self._set_cell_text(guest_table, 2, 1, guest_data.get('bairro', ''))
                self._set_cell_text(guest_table, 2, 5, guest_data.get('celular', guest_data.get('phone', '')))

                # Row 3: Cidade / Estado
                self._set_cell_text(guest_table, 3, 1, guest_data.get('cidade', ''))
                self._set_cell_text(guest_table, 3, 5, guest_data.get('estado', ''))

                # Row 4: Entrada / Saída / CEP
                self._set_cell_text(guest_table, 4, 1, check_in_formatted)
                self._set_cell_text(guest_table, 4, 3, check_out_formatted)
                self._set_cell_text(guest_table, 4, 5, guest_data.get('cep', ''))

            # ========== TABLE 2: Acompanhantes ==========
            companions = guest_data.get('companions', []) or []
            if len(doc.tables) >= 3 and companions:
                comp_table = doc.tables[2]
                for i, companion in enumerate(companions[:5]):  # Máximo 5 acompanhantes
                    if i < len(comp_table.rows):
                        name = companion.get('name', '') if isinstance(companion, dict) else str(companion)
                        doc_num = companion.get('document', '') if isinstance(companion, dict) else ''
                        self._set_cell_text(comp_table, i, 1, name)
                        self._set_cell_text(comp_table, i, 2, doc_num)

            # ========== VEÍCULO/PLACA (Parágrafo) ==========
            vehicle = guest_data.get('vehicle', '')
            plate = guest_data.get('plate', '')
            if vehicle or plate:
                for para in doc.paragraphs:
                    if 'culo/Modelo' in para.text or 'Veículo' in para.text or 'culo' in para.text:
                        vehicle_text = f"Veículo/Modelo: {vehicle}"
                        plate_text = f"Placa: {plate}"
                        para.text = f"{vehicle_text}    {plate_text}"
                        break

            # ========== DATA (Parágrafo com Brasília/DF) ==========
            date_today = datetime.now().strftime('%d/%m/%Y')
            for para in doc.paragraphs:
                if 'lia/DF' in para.text or 'Brasília' in para.text or 'Bras' in para.text:
                    para.text = f"Brasília/DF, {date_today}"
                    break

            # Salvar ou retornar bytes
            if save_to_file:
                filename = self._generate_filename(guest_data.get('name', 'hospede'))
                file_path = self.output_dir / filename
                doc.save(str(file_path))

                logger.info(f"Document generated: {file_path}")

                return {
                    "success": True,
                    "file_path": str(file_path),
                    "filename": filename,
                    "message": "Documento gerado com sucesso"
                }
            else:
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)

                return {
                    "success": True,
                    "file_bytes": doc_bytes.getvalue(),
                    "filename": self._generate_filename(guest_data.get('name', 'hospede')),
                    "message": "Documento gerado com sucesso"
                }

        except Exception as e:
            logger.error(f"Error generating document: {e}")
            return {
                "success": False,
                "message": "Erro ao gerar documento. Verifique se o template existe e tente novamente."
            }

    def _generate_filename(self, guest_name: str) -> str:
        """
        Gera nome de arquivo único para o documento.
        Formato: autorizacao_NOME_DDMMYYYY_HHMMSS.docx
        """
        clean_name = "".join(c for c in guest_name if c.isalnum() or c.isspace())
        clean_name = clean_name.replace(' ', '_')[:30]
        timestamp = datetime.now().strftime('%d%m%Y_%H%M%S')
        return f"autorizacao_{clean_name}_{timestamp}.docx"

    def _create_default_template(self, template_path: Path):
        """
        Cria um template padrão simples caso nenhum seja encontrado.
        """
        doc = Document()
        doc.add_heading('AUTORIZAÇÃO DE HOSPEDAGEM', 0)
        doc.add_paragraph('Template padrão - substitua pelo template real do condomínio.')

        # Criar tabela de hóspede (5x6)
        table = doc.add_table(rows=5, cols=6)
        table.style = 'Table Grid'
        labels = [
            ['Hóspede:', '', '', '', 'CPF:', ''],
            ['Endereço:', '', '', '', 'Telefone:', ''],
            ['Bairro:', '', '', '', 'Celular:', ''],
            ['Cidade:', '', '', '', 'Estado:', ''],
            ['Entrada:', '', 'Saída:', '', 'CEP:', ''],
        ]
        for ri, row_data in enumerate(labels):
            for ci, text in enumerate(row_data):
                table.rows[ri].cells[ci].text = text

        doc.save(str(template_path))
        logger.info(f"Default template created: {template_path}")

    def analyze_pdf_template(self, pdf_bytes: bytes) -> Dict[str, Any]:
        """
        Analisa um PDF template e detecta automaticamente campos de dados via pattern matching.

        Usa PyMuPDF para extrair blocos de texto com posições (bounding boxes).
        Para cada linha de texto, verifica se corresponde a um rótulo de campo conhecido.

        Campos detectados: guest_name, cpf, check_in, check_out, phone, celular,
                           address, bairro, cidade, estado, cep, vehicle, plate,
                           companion_1..5

        Salva o mapeamento em template_map.json no diretório de templates.

        Args:
            pdf_bytes: Conteúdo binário do arquivo PDF

        Returns:
            dict com 'fields' (mapeamento detectado), 'total_pages', 'created_at'
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.error("PyMuPDF não instalado. Execute: pip install PyMuPDF")
            return {"fields": {}, "total_pages": 0, "created_at": datetime.now().isoformat(), "error": "PyMuPDF não instalado"}

        # Padrões de rótulos para cada campo lógico (case-insensitive)
        FIELD_PATTERNS: Dict[str, List[str]] = {
            "guest_name":  [r"h[oó]spede\s*[:\-]?", r"nome\s+completo\s*[:\-]?", r"nome\s*[:\-]"],
            "cpf":         [r"c\.?p\.?f\.?\s*[:\-]?", r"cpf\s*[:\-]"],
            "check_in":    [r"check[\s\-]?in\s*[:\-]?", r"entrada\s*[:\-]?", r"chegada\s*[:\-]?"],
            "check_out":   [r"check[\s\-]?out\s*[:\-]?", r"sa[ií]da\s*[:\-]?", r"partida\s*[:\-]?"],
            "phone":       [r"telefone\s*[:\-]?", r"fone\s*[:\-]?", r"tel\.?\s*[:\-]?"],
            "celular":     [r"celular\s*[:\-]?", r"cel\.?\s*[:\-]?", r"whatsapp\s*[:\-]?"],
            "address":     [r"endere[cç]o\s*[:\-]?", r"rua\s*[:\-]?", r"av\.?\s*[:\-]?"],
            "bairro":      [r"bairro\s*[:\-]?"],
            "cidade":      [r"cidade\s*[:\-]?", r"munic[ií]pio\s*[:\-]?"],
            "estado":      [r"estado\s*[:\-]?", r"uf\s*[:\-]?"],
            "cep":         [r"c\.?e\.?p\.?\s*[:\-]?", r"cep\s*[:\-]?"],
            "vehicle":     [r"ve[ií]culo\s*[:\-]?", r"modelo\s*[:\-]?", r"carro\s*[:\-]?"],
            "plate":       [r"placa\s*[:\-]?"],
        }

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        fields: Dict[str, Any] = {}

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_dict = page.get_text("dict")

            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # Apenas blocos de texto
                    continue

                for line in block.get("lines", []):
                    # Concatenar todos os spans da linha em um texto único
                    line_text = " ".join(span.get("text", "") for span in line.get("spans", []))
                    line_lower = line_text.lower().strip()

                    if not line_lower:
                        continue

                    # Verificar cada campo
                    for field_name, patterns in FIELD_PATTERNS.items():
                        if field_name in fields:
                            continue  # Já detectado

                        for pattern in patterns:
                            if re.search(pattern, line_lower):
                                fields[field_name] = {
                                    "page": page_num,
                                    "bbox": list(line["bbox"]),  # [x0, y0, x1, y1]
                                    "detected_text": line_text.strip(),
                                }
                                logger.info(f"[PDF Analysis] Campo '{field_name}' detectado na página {page_num}: '{line_text.strip()}'")
                                break

        doc.close()

        # Montar resultado e salvar como template_map.json
        template_map = {
            "source": "pdf_analysis",
            "fields": fields,
            "total_pages": len(doc),
            "fields_detected": len(fields),
            "created_at": datetime.now().isoformat(),
        }

        map_path = self.template_dir / "template_map.json"
        map_path.write_text(json.dumps(template_map, indent=2, ensure_ascii=False), encoding="utf-8")
        logger.info(f"[PDF Analysis] Mapeamento salvo em {map_path}. Campos detectados: {list(fields.keys())}")

        return template_map

    def get_template_map(self) -> Optional[Dict[str, Any]]:
        """
        Retorna o mapeamento de campos do template personalizado, se existir.

        Returns:
            dict do template_map.json ou None se não houver mapeamento
        """
        map_path = self.template_dir / "template_map.json"
        if not map_path.exists():
            return None
        try:
            return json.loads(map_path.read_text(encoding="utf-8"))
        except Exception as e:
            logger.error(f"Erro ao ler template_map.json: {e}")
            return None

    def list_generated_documents(self, limit: int = 50, offset: int = 0) -> list[Dict[str, Any]]:
        """Lista os documentos gerados, com suporte a paginação por offset."""
        documents = []

        all_files = sorted(
            self.output_dir.glob("*.docx"),
            key=lambda p: p.stat().st_mtime,
            reverse=True
        )

        for file_path in all_files[offset : offset + limit]:
            stat = file_path.stat()
            documents.append({
                "filename": file_path.name,
                "path": str(file_path),
                "size_kb": round(stat.st_size / 1024, 2),
                "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            })

        return documents

    def get_document_path(self, filename: str) -> Optional[Path]:
        """Retorna o caminho completo de um documento gerado."""
        file_path = self.output_dir / filename
        return file_path if file_path.exists() else None

    def delete_document(self, filename: str) -> bool:
        """Deleta um documento gerado."""
        try:
            file_path = self.output_dir / filename
            if file_path.exists():
                file_path.unlink()
                logger.info(f"Document deleted: {filename}")
                return True
            return False
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False
